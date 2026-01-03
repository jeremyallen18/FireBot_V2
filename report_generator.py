from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime, timedelta
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from io import BytesIO
import json
import os
from config import LOCATION_INFO


class FireBotReportGenerator:
    """Generador de reportes Word profesionales para el sistema FireBot"""

    def __init__(self, data_file='firebot_history.json'):
        self.data_file = data_file
        self.history = self.load_history()

    # ==========================
    #   HISTORIAL
    # ==========================
    def load_history(self):
        if os.path.exists(self.data_file):
            try:
                with open(self.data_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return []
        return []

    def save_detection(self, detection_data, camera_frame=None):
        """
        Guarda una detección en el historial
        
        Args:
            detection_data: Diccionario con datos de la detección
            camera_frame: Frame de OpenCV (numpy array) de la cámara (opcional)
        """
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        img_path = None
        
        # Guardar captura de cámara si se proporciona
        if camera_frame is not None:
            try:
                import cv2
                detections_dir = 'detecciones'
                os.makedirs(detections_dir, exist_ok=True)
                
                relative_path = f"{detections_dir}/fire_{timestamp}.jpg"
                
                # Guardar con mayor calidad
                cv2.imwrite(relative_path, camera_frame, [cv2.IMWRITE_JPEG_QUALITY, 95])
                
                # Verificar que la imagen se guardó correctamente
                if os.path.exists(relative_path) and os.path.getsize(relative_path) > 0:
                    img_path = relative_path
                    print(f"Imagen guardada en: {os.path.abspath(relative_path)}")
                else:
                    print(f"Advertencia: La imagen no se guardó correctamente")
                    
            except Exception as e:
                print(f"Error al guardar imagen: {e}")
                import traceback
                traceback.print_exc()
        
        detection = {
            'timestamp': datetime.now().isoformat(),
            'confidence': detection_data.get('confidence', 0),
            'duration': detection_data.get('duration', 0),
            'alarm_triggered': detection_data.get('alarm_triggered', False),
            'manual_silence': detection_data.get('manual_silence', False),
            'image_path': img_path,
            'latitude': detection_data.get('latitude', LOCATION_INFO.get('latitude')),
            'longitude': detection_data.get('longitude', LOCATION_INFO.get('longitude')),
            'location_name': detection_data.get('location_name', LOCATION_INFO.get('zona'))
        }

        self.history.append(detection)

        with open(self.data_file, 'w', encoding='utf-8') as f:
            json.dump(self.history, f, indent=2, ensure_ascii=False)

    # ==========================
    #   FILTRO
    # ==========================
    def filter_by_period(self, period='week'):
        now = datetime.now()
        start = now - timedelta(days=7 if period == 'week' else 30)

        return [
            d for d in self.history
            if datetime.fromisoformat(d['timestamp']) >= start
        ]

    # ==========================
    #   GRÁFICAS
    # ==========================
    def create_chart(self, detections, chart_type='timeline'):
        if not detections:
            return None

        fig, ax = plt.subplots(figsize=(8, 4))
        fig.patch.set_facecolor('white')
        ax.set_facecolor('white')

        if chart_type == 'timeline':
            dates = [datetime.fromisoformat(d['timestamp']) for d in detections]
            conf = [d['confidence'] for d in detections]

            ax.plot(dates, conf, marker='o', color='#c0392b', linewidth=2, markersize=6)
            ax.fill_between(dates, conf, alpha=0.2, color='#e74c3c')
            ax.set_title('Nivel de Confianza de Detección por Fecha', 
                        fontweight='bold', fontsize=12, pad=15)
            ax.set_ylabel('Nivel de Confianza (%)', fontsize=10)
            ax.set_xlabel('Fecha', fontsize=10)
            ax.xaxis.set_major_formatter(mdates.DateFormatter('%d/%m'))
            ax.grid(True, alpha=0.3, linestyle='--')
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            plt.xticks(rotation=45)

        plt.tight_layout()

        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=200, bbox_inches='tight')
        buffer.seek(0)
        plt.close()
        return buffer

    # ==========================
    #   MAPA DE UBICACIÓN
    # ==========================
    def generate_map_image(self, location_info):
        """
        Genera una imagen mostrando la información de ubicación
        """
        fig, ax = plt.subplots(figsize=(7, 3), facecolor='white')

        ax.set_title("UBICACIÓN DEL SISTEMA", fontsize=13, fontweight='bold', pad=20)
        
        info_text = (
            f"Zona: {location_info.get('zona', 'N/A')}\n"
            f"Edificio: {location_info.get('edificio', 'N/A')}\n"
            f"Identificador de Cámara: {location_info.get('camara_id', 'N/A')}\n"
        )
        
        if 'latitude' in location_info and 'longitude' in location_info:
            info_text += f"\nCoordenadas Geográficas:\n{location_info['latitude']}, {location_info['longitude']}"
        
        ax.text(
            0.5, 0.5,
            info_text,
            fontsize=11,
            ha='center',
            va='center',
            bbox=dict(boxstyle='round,pad=1', facecolor='#ecf0f1', edgecolor='#34495e', linewidth=1.5)
        )

        ax.axis('off')

        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=200, bbox_inches='tight', facecolor='white')
        buffer.seek(0)
        plt.close()

        return buffer

    # ==========================
    #   ESTILOS DE DOCUMENTO
    # ==========================
    def setup_document_styles(self, doc):
        """Configura los estilos personalizados del documento"""
        
        # Estilo para el título principal
        styles = doc.styles
        
        # Título principal
        if 'ReportTitle' not in styles:
            title_style = styles.add_style('ReportTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Pt(24)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(44, 62, 80)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(30)
        
        # Subtítulo
        if 'ReportSubtitle' not in styles:
            subtitle_style = styles.add_style('ReportSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            subtitle_style.font.size = Pt(12)
            subtitle_style.font.color.rgb = RGBColor(127, 140, 141)
            subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_style.paragraph_format.space_after = Pt(40)

    # ==========================
    #   REPORTE WORD
    # ==========================
    def generate_report(self, period='week', filename=None):
        detections = self.filter_by_period(period)

        if filename is None:
            filename = f"FireBot_Reporte_{period}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        # Crear documento
        doc = Document()
        self.setup_document_styles(doc)
        
        # Configurar márgenes
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # ==========================
        #   PORTADA
        # ==========================
        title = doc.add_paragraph('FIREBOT', style='ReportTitle')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Sistema de Detección de Incendios con Inteligencia Artificial')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_format = subtitle.runs[0]
        subtitle_format.font.size = Pt(12)
        subtitle_format.font.color.rgb = RGBColor(127, 140, 141)
        
        doc.add_paragraph()  # Espacio
        
        # Información del reporte
        period_text = "Semanal" if period == 'week' else "Mensual"
        info = doc.add_paragraph()
        info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info.add_run(
            f'REPORTE {period_text.upper()}\n'
            f'Fecha de generación: {datetime.now().strftime("%d de %B de %Y")}\n'
            f'Hora: {datetime.now().strftime("%H:%M:%S")}'
        )
        info_run.font.size = Pt(11)
        
        doc.add_page_break()

        # ==========================
        #   TABLA DE CONTENIDOS
        # ==========================
        heading = doc.add_heading('CONTENIDO', level=1)
        heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
        
        toc_items = [
            '1. Información de Ubicación',
            '2. Resumen Estadístico',
            '3. Análisis de Tendencias',
            '4. Registro de Eventos'
        ]
        
        for item in toc_items:
            p = doc.add_paragraph(item, style='List Number')
            p.runs[0].font.size = Pt(11)
        
        doc.add_page_break()

        # ==========================
        #   1. INFORMACIÓN DE UBICACIÓN
        # ==========================
        heading = doc.add_heading('1. INFORMACIÓN DE UBICACIÓN', level=1)
        heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
        
        # Generar y agregar mapa
        map_img = self.generate_map_image(LOCATION_INFO)
        doc.add_picture(map_img, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espacio

        # ==========================
        #   2. RESUMEN ESTADÍSTICO
        # ==========================
        heading = doc.add_heading('2. RESUMEN ESTADÍSTICO', level=1)
        heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)

        stats = self._calculate_statistics(detections)
        
        # Período de análisis
        if detections:
            first_date = datetime.fromisoformat(detections[0]['timestamp']).strftime('%d/%m/%Y')
            last_date = datetime.fromisoformat(detections[-1]['timestamp']).strftime('%d/%m/%Y')
            period_info = f"Período analizado: {first_date} - {last_date}"
        else:
            period_info = f"Período analizado: {period_text}"
        
        p = doc.add_paragraph(period_info)
        p.runs[0].font.size = Pt(10)
        
        doc.add_paragraph()  # Espacio

        # Tabla de estadísticas
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Encabezados
        header_cells = table.rows[0].cells
        header_cells[0].text = 'MÉTRICA'
        header_cells[1].text = 'VALOR'
        
        # Formatear encabezados
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            # Aplicar color de fondo usando XML
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shading_elm = parse_xml(r'<w:shd {} w:fill="2C3E50"/>'.format(nsdecls('w')))
            cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Datos
        data_rows = [
            ['Total de Detecciones', str(stats['total_detections'])],
            ['Alarmas Activadas', str(stats['alarms_triggered'])],
            ['Nivel de Confianza Promedio', f"{stats['avg_confidence']:.1f}%"],
            ['Duración Total de Eventos', f"{stats['total_duration']:.1f} segundos"]
        ]
        
        for i, row_data in enumerate(data_rows, start=1):
            cells = table.rows[i].cells
            cells[0].text = row_data[0]
            cells[1].text = row_data[1]

        doc.add_paragraph()  # Espacio

        # ==========================
        #   3. ANÁLISIS DE TENDENCIAS
        # ==========================
        if detections:
            doc.add_page_break()
            heading = doc.add_heading('3. ANÁLISIS DE TENDENCIAS', level=1)
            heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)
            
            description = doc.add_paragraph(
                'El siguiente gráfico muestra la evolución del nivel de confianza de las detecciones '
                'a lo largo del período analizado. Los valores superiores al 60% indican detecciones '
                'de alta confiabilidad que requieren atención inmediata.'
            )
            description.runs[0].font.size = Pt(10)
            
            doc.add_paragraph()  # Espacio
            
            # Agregar gráfica
            chart = self.create_chart(detections)
            if chart:
                doc.add_picture(chart, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ==========================
        #   4. REGISTRO DE EVENTOS
        # ==========================
        if detections:
            doc.add_page_break()
            heading = doc.add_heading('4. REGISTRO DETALLADO DE EVENTOS', level=1)
            heading.runs[0].font.color.rgb = RGBColor(44, 62, 80)

            intro_text = doc.add_paragraph(
                f'A continuación se presenta el detalle de los últimos {min(5, len(detections))} eventos '
                f'registrados por el sistema, incluyendo información técnica y evidencia fotográfica cuando '
                f'está disponible.'
            )
            intro_text.runs[0].font.size = Pt(10)
            
            doc.add_paragraph()  # Espacio

            # Mostrar los últimos 5 eventos
            for i, det in enumerate(detections[-5:], 1):
                if i > 1:
                    doc.add_page_break()
                
                event_heading = doc.add_heading(f'Evento #{i}', level=2)
                event_heading.runs[0].font.color.rgb = RGBColor(52, 73, 94)

                dt = datetime.fromisoformat(det['timestamp']).strftime('%d/%m/%Y %H:%M:%S')
                location = det.get('location_name', LOCATION_INFO.get('zona', 'No especificada'))
                
                # Tabla de información del evento
                event_table = doc.add_table(rows=6, cols=2)
                event_table.style = 'Light Grid'
                
                event_data = [
                    ['FECHA Y HORA', dt],
                    ['UBICACIÓN', location],
                    ['NIVEL DE CONFIANZA', f"{det['confidence']:.1f}%"],
                    ['DURACIÓN DEL EVENTO', f"{det['duration']:.1f} segundos"],
                    ['ESTADO DE ALARMA', 'Activada' if det['alarm_triggered'] else 'No activada'],
                    ['SILENCIAMIENTO MANUAL', 'Sí' if det.get('manual_silence', False) else 'No'],
                ]
                
                for row_idx, (label, value) in enumerate(event_data):
                    cells = event_table.rows[row_idx].cells
                    cells[0].text = label
                    cells[1].text = value
                    cells[0].paragraphs[0].runs[0].font.bold = True
                    
                    # Color de fondo para la primera columna usando XML
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    shading_elm = parse_xml(r'<w:shd {} w:fill="ECF0F1"/>'.format(nsdecls('w')))
                    cells[0]._element.get_or_add_tcPr().append(shading_elm)

                doc.add_paragraph()  # Espacio

                # Imagen de evidencia
                img_path = det.get('image_path')
                if img_path:
                    # Manejar tanto rutas absolutas como relativas
                    if not os.path.isabs(img_path):
                        img_path = os.path.abspath(img_path)
                    
                    if os.path.exists(img_path):
                        evidence_heading = doc.add_heading('EVIDENCIA VISUAL', level=3)
                        evidence_heading.runs[0].font.color.rgb = RGBColor(52, 73, 94)
                        
                        try:
                            # Verificar que el archivo sea una imagen válida
                            from PIL import Image as PILImage
                            test_img = PILImage.open(img_path)
                            test_img.verify()
                            
                            # Agregar la imagen
                            doc.add_picture(img_path, width=Inches(5.5))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            print(f"Imagen añadida al reporte: {img_path}")
                        except Exception as e:
                            error_msg = doc.add_paragraph(f"Error al cargar la evidencia visual: {str(e)}")
                            error_msg.runs[0].font.italic = True
                            error_msg.runs[0].font.size = Pt(9)
                            print(f"Error al agregar imagen: {e}")
                            import traceback
                            traceback.print_exc()
                    else:
                        warning = doc.add_paragraph("Advertencia: Imagen no encontrada en la ruta especificada.")
                        warning.runs[0].font.italic = True
                        warning.runs[0].font.size = Pt(9)
                        print(f"Imagen no encontrada: {img_path}")
                else:
                    no_evidence = doc.add_paragraph("Evidencia visual no disponible para este evento.")
                    no_evidence.runs[0].font.italic = True
                    no_evidence.runs[0].font.size = Pt(9)

        else:
            doc.add_paragraph(
                "No se registraron eventos de detección durante el período analizado."
            )

        # ==========================
        #   PIE DE PÁGINA
        # ==========================
        doc.add_page_break()
        
        # Agregar línea divisoria
        footer_line = doc.add_paragraph('_' * 80)
        footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Espacio
        
        # Información final
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(
            f'FIREBOT - SISTEMA DE DETECCIÓN DE INCENDIOS\n'
            f'Sistema de Monitoreo con Inteligencia Artificial\n\n'
            f'Reporte generado el {datetime.now().strftime("%d de %B de %Y a las %H:%M:%S")}\n'
            f'Período de análisis: {period_text}\n\n'
            f'© 2024 FireBot D.A.T. - Todos los derechos reservados\n'
            f'Este documento es confidencial y está destinado únicamente para uso interno'
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(149, 165, 166)

        # Guardar el documento
        try:
            doc.save(filename)
            print(f"Reporte generado exitosamente: {filename}")
            return filename
        except Exception as e:
            print(f"Error al generar el documento Word: {e}")
            import traceback
            traceback.print_exc()
            raise

    # ==========================
    #   ESTADÍSTICAS
    # ==========================
    def _calculate_statistics(self, detections):
        if not detections:
            return {
                'total_detections': 0,
                'alarms_triggered': 0,
                'avg_confidence': 0,
                'total_duration': 0
            }

        return {
            'total_detections': len(detections),
            'alarms_triggered': sum(d['alarm_triggered'] for d in detections),
            'avg_confidence': sum(d['confidence'] for d in detections) / len(detections),
            'total_duration': sum(d['duration'] for d in detections)
        }